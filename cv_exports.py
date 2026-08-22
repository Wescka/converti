
from __future__ import annotations

import base64
import io
import re
import zipfile
from xml.sax.saxutils import escape as xml_escape
from typing import Any

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageTemplate,
    Paragraph,
    Spacer,
)

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


LABELS = {
    "es": {
        "profile":"Perfil profesional","experience":"Experiencia laboral","education":"Educación",
        "skills":"Habilidades","languages":"Idiomas","certifications":"Certificaciones","contact":"Contacto"
    },
    "en": {
        "profile":"Professional profile","experience":"Work experience","education":"Education",
        "skills":"Skills","languages":"Languages","certifications":"Certifications","contact":"Contact"
    },
    "fr": {
        "profile":"Profil professionnel","experience":"Expérience professionnelle","education":"Formation",
        "skills":"Compétences","languages":"Langues","certifications":"Certifications","contact":"Contact"
    },
    "pt-br": {
        "profile":"Perfil profissional","experience":"Experiência profissional","education":"Educação",
        "skills":"Habilidades","languages":"Idiomas","certifications":"Certificações","contact":"Contato"
    },
}


def _hex(value: str, default="#2a7bff") -> str:
    raw = re.sub(r"[^0-9A-Fa-f]", "", str(value or ""))
    if len(raw) != 6:
        raw = default.lstrip("#")
    return "#" + raw.lower()


def _rgb(value: str) -> RGBColor:
    value = _hex(value).lstrip("#")
    return RGBColor(int(value[0:2],16), int(value[2:4],16), int(value[4:6],16))


def _safe(v: Any) -> str:
    return str(v or "").strip()

def _pdf_text(v: Any) -> str:
    """Escape user text before passing it to ReportLab Paragraph markup."""
    return xml_escape(_safe(v), {"\"": "&quot;"})


def _photo_bytes(data_url: str) -> bytes | None:
    if not isinstance(data_url, str) or not data_url.startswith("data:image/") or "," not in data_url:
        return None
    try:
        raw = base64.b64decode(data_url.split(",",1)[1], validate=False)
        return raw if len(raw) <= 8 * 1024 * 1024 else None
    except Exception:
        return None


def _initials(name: str) -> str:
    parts = [x for x in re.split(r"\s+", _safe(name)) if x]
    return "".join(x[0] for x in parts[:2]).upper() or "CV"


# ---------------------------- PDF ---------------------------------

def _pdf_styles(accent: str):
    base = getSampleStyleSheet()
    ink = colors.HexColor("#122447")
    muted = colors.HexColor("#667085")
    acc = colors.HexColor(accent)
    return {
        "name": ParagraphStyle("CVName", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=20, leading=21, textColor=ink, spaceAfter=2),
        "title": ParagraphStyle("CVTitle", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=9.5, leading=12, textColor=acc, spaceAfter=8),
        "section": ParagraphStyle("CVSection", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=8.5, leading=10, textColor=acc, spaceBefore=6, spaceAfter=4),
        "body": ParagraphStyle("CVBody", parent=base["Normal"], fontName="Helvetica",
            fontSize=8.2, leading=11.1, textColor=colors.HexColor("#475467"), spaceAfter=4),
        "entry": ParagraphStyle("CVEntry", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=8.4, leading=10.5, textColor=ink, spaceAfter=1),
        "meta": ParagraphStyle("CVMeta", parent=base["Normal"], fontName="Helvetica",
            fontSize=7.3, leading=9, textColor=muted, spaceAfter=2),
        "side_h": ParagraphStyle("CVSideH", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=8, leading=10, textColor=colors.white, spaceBefore=7, spaceAfter=4),
        "side": ParagraphStyle("CVSide", parent=base["Normal"], fontName="Helvetica",
            fontSize=7.2, leading=9.2, textColor=colors.HexColor("#e5edf8"), spaceAfter=2),
        "side_chip": ParagraphStyle("CVSideChip", parent=base["Normal"], fontName="Helvetica",
            fontSize=6.8, leading=8.2, textColor=colors.white, spaceAfter=2),
    }


def generate_cv_pdf(cv: dict, accent="#2a7bff", template="modern", locale="es", photo_data="") -> bytes:
    """Generate a complete multi-page PDF. Content flows; it is never clipped at 297 mm."""
    accent = _hex(accent)
    template = _safe(template).lower() or "modern"
    labels = LABELS.get(locale, LABELS["es"])
    styles = _pdf_styles(accent)
    out = io.BytesIO()

    modern = template == "modern"
    if modern:
        left_w = 57 * mm
        main_left = 68 * mm
        main_right = 15 * mm
        main_bottom = 14 * mm
        main_top = 15 * mm
        frame = Frame(main_left, main_bottom, A4[0]-main_left-main_right,
                      A4[1]-main_bottom-main_top, id="main", leftPadding=0, rightPadding=0,
                      topPadding=0, bottomPadding=0)
    else:
        margin = 17 * mm
        frame = Frame(margin, margin, A4[0]-2*margin, A4[1]-2*margin, id="main",
                      leftPadding=0,rightPadding=0,topPadding=0,bottomPadding=0)

    photo = _photo_bytes(photo_data)
    def draw_sidebar(canvas, doc):
        canvas.saveState()
        page_w, page_h = A4
        if modern:
            canvas.setFillColor(colors.HexColor("#13294b"))
            canvas.rect(0, 0, 57*mm, page_h, stroke=0, fill=1)
            x = 8*mm
            y = page_h - 18*mm

            # Photo/initials circle.
            cx, cy, radius = 28.5*mm, y-12*mm, 11.5*mm
            canvas.setStrokeColor(colors.HexColor("#d9e4f2"))
            canvas.setLineWidth(1)
            canvas.circle(cx, cy, radius, stroke=1, fill=0)
            if photo:
                try:
                    img = ImageReader(io.BytesIO(photo))
                    canvas.saveState()
                    p = canvas.beginPath()
                    p.circle(cx, cy, radius)
                    canvas.clipPath(p, stroke=0, fill=0)
                    canvas.drawImage(img, cx-radius, cy-radius, 2*radius, 2*radius,
                                     preserveAspectRatio=True, anchor="c", mask="auto")
                    canvas.restoreState()
                except Exception:
                    photo_fallback = True
                else:
                    photo_fallback = False
            else:
                photo_fallback = True
            if photo_fallback:
                canvas.setFillColor(colors.white)
                canvas.setFont("Helvetica-Bold", 10)
                canvas.drawCentredString(cx, cy-3, _initials(cv.get("name","")))

            text_y = cy - radius - 12*mm
            def side_heading(txt):
                nonlocal text_y
                canvas.setFillColor(colors.white)
                canvas.setFont("Helvetica-Bold", 7.5)
                canvas.drawString(x, text_y, txt.upper())
                text_y -= 5.3*mm
            def side_line(txt, size=6.5):
                nonlocal text_y
                if not txt:
                    return
                canvas.setFillColor(colors.HexColor("#e5edf8"))
                canvas.setFont("Helvetica", size)
                # Simple fit-to-width.
                maxw = 41*mm
                s = txt
                while canvas.stringWidth(s, "Helvetica", size) > maxw and len(s) > 5:
                    s = s[:-2] + "…"
                canvas.drawString(x, text_y, s)
                text_y -= 4.2*mm

            side_heading(labels["contact"])
            for item in (cv.get("email"),cv.get("phone"),cv.get("city"),cv.get("website")):
                side_line(_safe(item))
            skills = [x.get("name","") for x in cv.get("skills",[]) if isinstance(x,dict) and x.get("name")][:8]
            if skills:
                text_y -= 2*mm; side_heading(labels["skills"])
                for item in skills:
                    side_line("• " + _safe(item), 6.2)
            langs = [x for x in cv.get("languages",[]) if isinstance(x,dict) and x.get("name")]
            if langs:
                text_y -= 2*mm; side_heading(labels["languages"])
                for item in langs[:8]:
                    side_line(f"{_safe(item.get('name'))}  {_safe(item.get('level'))}", 6.2)

        # Page number tiny / non-distracting.
        canvas.setFillColor(colors.HexColor("#98a2b3"))
        canvas.setFont("Helvetica", 6.5)
        canvas.drawRightString(A4[0]-8*mm, 6*mm, str(doc.page))
        canvas.restoreState()

    doc = BaseDocTemplate(out, pagesize=A4, leftMargin=0,rightMargin=0,topMargin=0,bottomMargin=0,
                          allowSplitting=1, pageCompression=1)
    doc.addPageTemplates([PageTemplate(id="cv", frames=[frame], onPage=draw_sidebar)])

    story = [
        Paragraph(_pdf_text(cv.get("name")) or "CV", styles["name"]),
        Paragraph(_pdf_text(cv.get("title")), styles["title"]),
        Spacer(1, 1.5*mm),
    ]
    if not modern:
        contact = " · ".join(x for x in [_safe(cv.get("email")),_safe(cv.get("phone")),_safe(cv.get("city")),_safe(cv.get("website"))] if x)
        if contact:
            story.append(Paragraph(_pdf_text(contact), styles["meta"]))
            story.append(Spacer(1, 2*mm))

    def heading(key):
        story.append(Paragraph(labels[key].upper(), styles["section"]))

    if _safe(cv.get("profile")):
        heading("profile")
        story.append(Paragraph(_pdf_text(cv.get("profile")), styles["body"]))

    if cv.get("experience"):
        heading("experience")
        for item in cv["experience"]:
            if not isinstance(item,dict): continue
            role, company, period, desc = map(_safe, (item.get("role"),item.get("company"),item.get("period"),item.get("description")))
            head = role or company
            entry_flow = []
            if head:
                parts = [f"<b>{_pdf_text(head)}</b>"]
                if company and company != head: parts.append(f" · {_pdf_text(company)}")
                if period: parts.append(f"   <font color='#667085'>{_pdf_text(period)}</font>")
                entry_flow.append(Paragraph("".join(parts), styles["entry"]))
            if desc:
                entry_flow.append(Paragraph(_pdf_text(desc), styles["body"]))
            entry_flow.append(Spacer(1, 1.3*mm))
            if entry_flow:
                story.append(KeepTogether(entry_flow))

    if cv.get("education"):
        heading("education")
        for item in cv["education"]:
            if not isinstance(item,dict): continue
            degree, school, period, desc = map(_safe, (item.get("degree"),item.get("school"),item.get("period"),item.get("description")))
            line = degree or school
            if line:
                parts=[f"<b>{_pdf_text(line)}</b>"]
                if school and school != line: parts.append(f" · {_pdf_text(school)}")
                if period: parts.append(f"   <font color='#667085'>{_pdf_text(period)}</font>")
                story.append(Paragraph("".join(parts), styles["entry"]))
            if desc: story.append(Paragraph(_pdf_text(desc), styles["body"]))

    if not modern and cv.get("skills"):
        heading("skills")
        story.append(Paragraph(_pdf_text(" · ".join(_safe(x.get("name")) for x in cv["skills"] if isinstance(x,dict) and x.get("name"))[:1600]), styles["body"]))
    if not modern and cv.get("languages"):
        heading("languages")
        story.append(Paragraph(_pdf_text(" · ".join((_safe(x.get("name")) + (f" — {_safe(x.get('level'))}" if x.get("level") else "")) for x in cv["languages"] if isinstance(x,dict) and x.get("name"))), styles["body"]))

    if cv.get("certifications"):
        heading("certifications")
        for item in cv["certifications"]:
            if not isinstance(item,dict): continue
            line = " · ".join(x for x in [_safe(item.get("name")),_safe(item.get("issuer")),_safe(item.get("year"))] if x)
            if line: story.append(Paragraph(_pdf_text(line), styles["body"]))

    doc.build(story)
    return out.getvalue()


# ---------------------------- DOCX --------------------------------

def _set_cell_shading(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#").upper())


def _set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def _remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top","left","bottom","right","insideH","insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def _paragraph(container, text="", size=9, bold=False, color="#475467", space_after=2, align=None):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(_safe(text))
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = _rgb(color)
    return p


def _docx_heading(container, text, accent, side=False):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(_safe(text).upper())
    r.font.name = "Aptos"
    r.font.size = Pt(8.5 if not side else 8)
    r.bold = True
    r.font.color.rgb = RGBColor(255,255,255) if side else _rgb(accent)
    return p


def _main_sections(container, cv, accent, labels):
    if _safe(cv.get("profile")):
        _docx_heading(container, labels["profile"], accent)
        _paragraph(container, cv.get("profile"), 9, color="#475467", space_after=4)
    if cv.get("experience"):
        _docx_heading(container, labels["experience"], accent)
        for item in cv["experience"]:
            if not isinstance(item,dict): continue
            role, company, period, desc = map(_safe, (item.get("role"),item.get("company"),item.get("period"),item.get("description")))
            p=container.add_paragraph()
            p.paragraph_format.space_after=Pt(5)
            p.paragraph_format.keep_together=True
            p.paragraph_format.widow_control=True
            r=p.add_run(role or company); r.bold=True; r.font.name="Aptos"; r.font.size=Pt(9); r.font.color.rgb=_rgb("#122447")
            if period:
                rr=p.add_run("    "+period); rr.font.name="Aptos"; rr.font.size=Pt(7.8); rr.font.color.rgb=_rgb("#7b879d")
            if company and company != (role or company):
                rr=p.add_run("\n"+company); rr.bold=True; rr.font.name="Aptos"; rr.font.size=Pt(8); rr.font.color.rgb=_rgb(accent)
            if desc:
                rr=p.add_run("\n"+desc); rr.font.name="Aptos"; rr.font.size=Pt(8.4); rr.font.color.rgb=_rgb("#536078")
    if cv.get("education"):
        _docx_heading(container, labels["education"], accent)
        for item in cv["education"]:
            if not isinstance(item,dict): continue
            line=" · ".join(x for x in [_safe(item.get("degree")),_safe(item.get("school")),_safe(item.get("period"))] if x)
            if line: _paragraph(container,line,8.6,bold=True,color="#122447",space_after=1)
            if _safe(item.get("description")): _paragraph(container,item.get("description"),8.2,color="#536078",space_after=3)
    if cv.get("certifications"):
        _docx_heading(container, labels["certifications"], accent)
        for item in cv["certifications"]:
            if not isinstance(item,dict): continue
            line=" · ".join(x for x in [_safe(item.get("name")),_safe(item.get("issuer")),_safe(item.get("year"))] if x)
            if line: _paragraph(container,line,8.4,color="#475467",space_after=2)


def generate_cv_docx(cv: dict, accent="#2a7bff", template="modern", locale="es", photo_data="") -> bytes:
    accent = _hex(accent)
    template = _safe(template).lower() or "modern"
    labels = LABELS.get(locale, LABELS["es"])
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
    sec.top_margin = sec.bottom_margin = Inches(0.42)
    sec.left_margin = sec.right_margin = Inches(0.42)
    sec.header_distance = sec.footer_distance = Inches(0.2)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)

    photo = _photo_bytes(photo_data)

    if template == "modern":
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(2.15)
        table.columns[1].width = Inches(5.25)
        _remove_table_borders(table)
        left, right = table.cell(0,0), table.cell(0,1)
        left.width, right.width = Inches(2.15), Inches(5.25)
        left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _set_cell_shading(left, "#13294b")
        _set_cell_margins(left, 180, 180, 180, 180)
        _set_cell_margins(right, 170, 300, 170, 160)

        # Remove empty initial paragraphs only visually by reusing them.
        p = left.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
        if photo:
            try: p.add_run().add_picture(io.BytesIO(photo), width=Inches(0.95))
            except Exception: p.add_run(_initials(cv.get("name","")))
        else:
            r=p.add_run(_initials(cv.get("name",""))); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor(255,255,255)

        _docx_heading(left, labels["contact"], accent, side=True)
        for item in (cv.get("email"),cv.get("phone"),cv.get("city"),cv.get("website")):
            if _safe(item): _paragraph(left,item,7.5,color="#e6eef9",space_after=2)
        skills=[x.get("name","") for x in cv.get("skills",[]) if isinstance(x,dict) and x.get("name")][:8]
        if skills:
            _docx_heading(left, labels["skills"], accent, side=True)
            for item in skills: _paragraph(left,"• "+item,7.3,color="#ffffff",space_after=1)
        langs=[x for x in cv.get("languages",[]) if isinstance(x,dict) and x.get("name")]
        if langs:
            _docx_heading(left, labels["languages"], accent, side=True)
            for x in langs[:8]:
                _paragraph(left,f"{_safe(x.get('name'))}  {_safe(x.get('level'))}",7.3,color="#ffffff",space_after=1)

        # Reuse first right paragraph for heading.
        p=right.paragraphs[0]; p.paragraph_format.space_after=Pt(1)
        r=p.add_run(_safe(cv.get("name")) or "CV"); r.bold=True; r.font.name="Aptos"; r.font.size=Pt(21); r.font.color.rgb=_rgb("#122447")
        if _safe(cv.get("title")): _paragraph(right,cv.get("title"),9.5,bold=True,color=accent,space_after=7)
        _main_sections(right,cv,accent,labels)
    else:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after=Pt(2)
        r=p.add_run(_safe(cv.get("name")) or "CV"); r.bold=True; r.font.name="Aptos"; r.font.size=Pt(22); r.font.color.rgb=_rgb("#122447")
        if _safe(cv.get("title")): _paragraph(doc,cv.get("title"),10,bold=True,color=accent,space_after=3)
        contact=" · ".join(x for x in [_safe(cv.get("email")),_safe(cv.get("phone")),_safe(cv.get("city")),_safe(cv.get("website"))] if x)
        if contact: _paragraph(doc,contact,8,color="#667085",space_after=6)
        _main_sections(doc,cv,accent,labels)
        if cv.get("skills"):
            _docx_heading(doc, labels["skills"], accent)
            _paragraph(doc," · ".join(_safe(x.get("name")) for x in cv["skills"] if isinstance(x,dict) and x.get("name")),8.5,color="#475467")
        if cv.get("languages"):
            _docx_heading(doc, labels["languages"], accent)
            _paragraph(doc," · ".join((_safe(x.get("name")) + (f" — {_safe(x.get('level'))}" if x.get("level") else "")) for x in cv["languages"] if isinstance(x,dict) and x.get("name")),8.5,color="#475467")

    bio=io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def validate_cv_docx_bytes(payload: bytes) -> None:
    """Reject corrupt or non-DOCX export payloads before they are sent to the user."""
    if not payload or not payload.startswith(b"PK"):
        raise ValueError("El archivo Word generado no tiene una estructura DOCX válida.")
    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as zf:
            names = set(zf.namelist())
            required = {"[Content_Types].xml", "word/document.xml", "_rels/.rels"}
            if not required.issubset(names):
                raise ValueError("El archivo Word generado está incompleto.")
            document_xml = zf.read("word/document.xml")
            if b"<w:document" not in document_xml:
                raise ValueError("El documento Word generado no contiene un documento editable válido.")
    except zipfile.BadZipFile as exc:
        raise ValueError("El archivo Word generado está corrupto.") from exc

def validate_cv_pdf_bytes(payload: bytes) -> None:
    """Basic structural validation that catches empty/corrupt PDF exports."""
    if not payload or not payload.startswith(b"%PDF-") or b"%%EOF" not in payload[-4096:]:
        raise ValueError("El PDF generado no tiene una estructura válida.")
